from __future__ import annotations

import json
import tempfile
import unittest
from pathlib import Path

from docx import Document
from pypdf import PdfWriter
from pypdf.generic import DecodedStreamObject, DictionaryObject, NameObject

from dataforge.application import DataForge
from dataforge.config import Settings


def write_text_pdf(path: Path, text: str) -> None:
    writer = PdfWriter()
    page = writer.add_blank_page(width=420, height=595)
    font = DictionaryObject(
        {
            NameObject("/Type"): NameObject("/Font"),
            NameObject("/Subtype"): NameObject("/Type1"),
            NameObject("/BaseFont"): NameObject("/Helvetica"),
        }
    )
    font_ref = writer._add_object(font)
    page[NameObject("/Resources")] = DictionaryObject(
        {NameObject("/Font"): DictionaryObject({NameObject("/F1"): font_ref})}
    )
    content = DecodedStreamObject()
    content.set_data(f"BT /F1 12 Tf 50 520 Td ({text}) Tj ET".encode("ascii"))
    page[NameObject("/Contents")] = writer._add_object(content)
    with path.open("wb") as handle:
        writer.write(handle)


class PrioritySourceFormatsTest(unittest.TestCase):
    def setUp(self):
        self.temporary = tempfile.TemporaryDirectory()
        self.root = Path(self.temporary.name)
        self.app = DataForge(Settings.load(self.root))

    def tearDown(self):
        self.temporary.cleanup()

    def test_priority_formats_reach_published_assets(self):
        files: list[tuple[Path, str]] = []

        text_file = self.root / "follow-up.txt"
        text_file.write_text("患者血压稳定，继续随访。", encoding="utf-8")
        files.append((text_file, "患者血压稳定"))

        markdown_file = self.root / "guide.md"
        markdown_file.write_text("# 用药指南\n\n每日按时服药。", encoding="utf-8")
        files.append((markdown_file, "每日按时服药"))

        csv_file = self.root / "faq.csv"
        csv_file.write_text("question,answer\n如何预约,通过医院小程序预约\n", encoding="utf-8")
        files.append((csv_file, "通过医院小程序预约"))

        word_file = self.root / "report.docx"
        document = Document()
        document.add_heading("出院记录", level=1)
        document.add_paragraph("患者恢复良好，可以出院。")
        table = document.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "复诊"
        table.cell(0, 1).text = "两周后"
        document.save(word_file)
        files.append((word_file, "患者恢复良好"))

        pdf_file = self.root / "english-report.pdf"
        write_text_pdf(pdf_file, "Medical follow-up is stable")
        files.append((pdf_file, "Medical follow-up is stable"))

        for source_file, expected_text in files:
            with self.subTest(suffix=source_file.suffix):
                result = self.app.flow(
                    source_file,
                    name=source_file.stem,
                    engine_override="native",
                )
                self.assertEqual(result.run["status"], "completed")
                self.assertEqual(result.asset_version["status"], "published")
                blob = self.app.blobs.resolve(result.asset_version["blob_uri"])
                records = [json.loads(line) for line in blob.read_text(encoding="utf-8").splitlines()]
                self.assertTrue(any(expected_text in item["content"] for item in records))


if __name__ == "__main__":
    unittest.main()

from __future__ import annotations

import csv
import json
import mimetypes
from pathlib import Path
from typing import Any, Iterable

from .blobs import BlobStore
from .database import MetadataStore
from .errors import ValidationError
from .models import IngestResult


SUPPORTED_SUFFIXES = {".pdf", ".csv", ".md", ".docx", ".txt", ".json", ".jsonl"}
CONTENT_FIELDS = ("raw_content", "content", "text", "body", "answer")


class SourceService:
    def __init__(self, store: MetadataStore, blobs: BlobStore):
        self.store = store
        self.blobs = blobs

    def ingest(
        self,
        file_path: str | Path,
        *,
        source_id: str | None = None,
        name: str | None = None,
        kind: str = "file",
        metadata: dict[str, Any] | None = None,
    ) -> IngestResult:
        path = Path(file_path).expanduser().resolve()
        if not path.is_file():
            raise ValidationError(f"Source file does not exist: {path}")
        if path.suffix.lower() not in SUPPORTED_SUFFIXES:
            supported = ", ".join(sorted(SUPPORTED_SUFFIXES))
            raise ValidationError(f"Unsupported source type {path.suffix!r}; supported: {supported}")

        blob_uri, sha256, size_bytes = self.blobs.put_file(path)
        source = (
            self.store.get_source(source_id)
            if source_id
            else self.store.create_source(name or path.stem, kind, metadata or {})
        )
        existing = self.store.find_source_version_by_hash(source["id"], sha256)
        if existing:
            return IngestResult(source=source, source_version=existing, created=False)

        media_type = mimetypes.guess_type(path.name)[0] or "application/octet-stream"
        version = self.store.create_source_version(
            source_id=source["id"],
            blob_uri=blob_uri,
            sha256=sha256,
            size_bytes=size_bytes,
            media_type=media_type,
            original_filename=path.name,
        )
        return IngestResult(source=source, source_version=version, created=True)


def materialize_source_records(
    source_file: Path,
    source_version: dict[str, Any],
    destination: Path,
) -> int:
    suffix = Path(source_version["original_filename"]).suffix.lower()
    destination.parent.mkdir(parents=True, exist_ok=True)
    count = 0
    with destination.open("w", encoding="utf-8") as handle:
        for index, record in enumerate(_read_records(source_file, suffix)):
            content = _extract_content(record)
            if not content.strip():
                continue
            materialized = {
                "document_id": f"{source_version['id']}:{index}",
                "source_id": source_version["source_id"],
                "source_version_id": source_version["id"],
                "source_record_index": index,
                "raw_content": content,
            }
            handle.write(json.dumps(materialized, ensure_ascii=False) + "\n")
            count += 1

    if count == 0:
        raise ValidationError("Source records do not contain any text")
    return count


def _read_records(source_file: Path, suffix: str) -> Iterable[Any]:
    if suffix in {".txt", ".md"}:
        yield source_file.read_text(encoding="utf-8-sig", errors="replace")
        return
    if suffix == ".pdf":
        yield from _read_pdf(source_file)
        return
    if suffix == ".docx":
        yield from _read_docx(source_file)
        return
    if suffix == ".json":
        payload = json.loads(source_file.read_text(encoding="utf-8-sig"))
        if isinstance(payload, list):
            yield from payload
        else:
            yield payload
        return
    if suffix == ".jsonl":
        with source_file.open(encoding="utf-8-sig") as handle:
            for line_number, line in enumerate(handle, start=1):
                if not line.strip():
                    continue
                try:
                    yield json.loads(line)
                except json.JSONDecodeError as exc:
                    raise ValidationError(f"Invalid JSONL at line {line_number}: {exc}") from exc
        return
    if suffix == ".csv":
        with source_file.open(encoding="utf-8-sig", newline="") as handle:
            yield from csv.DictReader(handle)
        return
    raise ValidationError(f"No reader for source suffix: {suffix}")


def _read_pdf(source_file: Path) -> Iterable[str]:
    try:
        from pypdf import PdfReader

        reader = PdfReader(source_file)
        if reader.is_encrypted:
            try:
                reader.decrypt("")
            except Exception as exc:
                raise ValidationError("PDF 文件已加密，无法读取") from exc
        for page_number, page in enumerate(reader.pages, start=1):
            try:
                text = page.extract_text() or ""
            except Exception as exc:
                raise ValidationError(f"无法读取 PDF 第 {page_number} 页：{exc}") from exc
            if text.strip():
                yield text
    except ValidationError:
        raise
    except Exception as exc:
        raise ValidationError(f"无法解析 PDF 文件：{exc}") from exc


def _read_docx(source_file: Path) -> Iterable[str]:
    try:
        from docx import Document

        document = Document(source_file)
    except Exception as exc:
        raise ValidationError(f"无法解析 Word 文件：{exc}") from exc

    blocks: list[str] = []
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            blocks.append(paragraph.text.strip())
    for table in document.tables:
        for row in table.rows:
            values = [cell.text.strip() for cell in row.cells]
            if any(values):
                blocks.append("\t".join(values))
    if blocks:
        yield "\n\n".join(blocks)


def _extract_content(record: Any) -> str:
    if isinstance(record, str):
        return record
    if isinstance(record, dict):
        for field in CONTENT_FIELDS:
            value = record.get(field)
            if isinstance(value, str):
                return value
        return json.dumps(record, ensure_ascii=False, sort_keys=True)
    return str(record)
